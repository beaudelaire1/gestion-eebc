from datetime import date
from io import BytesIO

import pytest
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse


PDF_BYTES = b"%PDF-1.4\n1 0 obj<</Type/Catalog>>endobj\ntrailer<</Root 1 0 R>>\n%%EOF"


def _create_document(*, admin_user, settings, tmp_path, name, content, content_type):
    from apps.documents.models import Document

    settings.MEDIA_ROOT = tmp_path
    upload = SimpleUploadedFile(name, content, content_type=content_type)
    return Document.objects.create(
        title=f"Document {name}",
        file=upload,
        file_name=name,
        file_size=len(content),
        file_type=content_type,
        uploaded_by=admin_user,
        visibility=Document.Visibility.STAFF,
    )


@pytest.mark.django_db
def test_document_stream_pdf_allows_browser_inline_viewer(authenticated_client, admin_user, settings, tmp_path):
    document = _create_document(
        admin_user=admin_user,
        settings=settings,
        tmp_path=tmp_path,
        name='notice.pdf',
        content=PDF_BYTES,
        content_type='application/pdf',
    )

    response = authenticated_client.get(reverse('documents:stream', args=[document.pk]))

    assert response.status_code == 200
    assert response.headers['Content-Type'].startswith('application/pdf')
    assert response.headers['Content-Disposition'].startswith('inline;')
    assert 'X-Frame-Options' not in response.headers


@pytest.mark.django_db
def test_document_stream_pdf_extension_allows_inline_even_with_generic_mime(authenticated_client, admin_user, settings, tmp_path):
    document = _create_document(
        admin_user=admin_user,
        settings=settings,
        tmp_path=tmp_path,
        name='notice.pdf',
        content=PDF_BYTES,
        content_type='application/octet-stream',
    )

    response = authenticated_client.get(reverse('documents:stream', args=[document.pk]))

    assert response.status_code == 200
    assert response.headers['Content-Disposition'].startswith('inline;')
    assert 'X-Frame-Options' not in response.headers
    assert document.is_previewable is True
    # Content-Type must be application/pdf regardless of stored file_type (octet-stream),
    # because SECURE_CONTENT_TYPE_NOSNIFF would otherwise cause the browser to download the file.
    assert response.headers['Content-Type'].startswith('application/pdf')


@pytest.mark.django_db
def test_document_preview_html_stays_sameorigin(authenticated_client, admin_user, settings, tmp_path):
    document = _create_document(
        admin_user=admin_user,
        settings=settings,
        tmp_path=tmp_path,
        name='notes.txt',
        content=b'Premiere ligne\nDeuxieme ligne',
        content_type='text/plain',
    )

    response = authenticated_client.get(reverse('documents:preview', args=[document.pk]))

    assert response.status_code == 200
    assert response.headers['X-Frame-Options'] == 'SAMEORIGIN'


@pytest.mark.django_db
def test_document_preview_docx_renders_content(authenticated_client, admin_user, settings, tmp_path):
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_heading('Compte rendu', level=1)
    docx.add_paragraph('Contenu du document Word.')
    buffer = BytesIO()
    docx.save(buffer)

    document = _create_document(
        admin_user=admin_user,
        settings=settings,
        tmp_path=tmp_path,
        name='compte-rendu.docx',
        content=buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )

    response = authenticated_client.get(reverse('documents:preview', args=[document.pk]))

    assert response.status_code == 200
    assert 'Compte rendu'.encode() in response.content
    assert b'Word (.docx)' in response.content


@pytest.mark.django_db
def test_document_preview_legacy_doc_shows_clear_message(authenticated_client, admin_user, settings, tmp_path):
    document = _create_document(
        admin_user=admin_user,
        settings=settings,
        tmp_path=tmp_path,
        name='ancien.doc',
        content=b'D0CF11E0 legacy doc placeholder',
        content_type='application/msword',
    )

    response = authenticated_client.get(reverse('documents:preview', args=[document.pk]))

    assert response.status_code == 200
    assert document.is_previewable is True
    assert b'Word 97-2003 (.doc)' in response.content
    assert 'anciens fichiers'.encode() in response.content


@pytest.mark.django_db
def test_generated_pdf_allows_browser_inline_viewer(authenticated_client, admin_user, monkeypatch):
    from apps.documents.models import GeneratedDocument

    monkeypatch.setattr('apps.documents.generation_views.render_generated_document_pdf', lambda doc: PDF_BYTES)

    document = GeneratedDocument.objects.create(
        title='Courrier de test',
        kind=GeneratedDocument.Kind.COURRIER,
        document_date=date(2026, 4, 18),
        body_html='<p>Contenu.</p>',
        created_by=admin_user,
    )

    response = authenticated_client.get(reverse('documents:generated_pdf', args=[document.pk]))

    assert response.status_code == 200
    assert response.headers['Content-Type'].startswith('application/pdf')
    assert response.headers['Content-Disposition'].startswith('inline;')
    assert 'X-Frame-Options' not in response.headers